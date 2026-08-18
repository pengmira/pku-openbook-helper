import argparse
import json
import re
from pathlib import Path

from docx import Document


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def iter_docx_text(path: Path):
    doc = Document(path)

    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize(cell.text) for cell in row.cells]
            text = normalize(" | ".join(cell for cell in cells if cell))
            if text:
                yield text


def extract(args):
    source = Path(args.source)
    output = Path(args.output)
    chunks = list(iter_docx_text(source))
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(chunks), encoding="utf-8")

    meta = {
        "source": str(source),
        "output": str(output),
        "chunks": len(chunks),
        "characters": sum(len(chunk) for chunk in chunks),
    }
    print(json.dumps(meta, ensure_ascii=False, indent=2))


def score_chunk(chunk: str, terms):
    score = 0
    for term in terms:
        if not term:
            continue
        score += chunk.count(term) * max(2, len(term))
    return score


def search(args):
    text_path = Path(args.text)
    chunks = [line.strip() for line in text_path.read_text(encoding="utf-8").splitlines() if line.strip()]
    terms = [term.strip() for term in re.split(r"[\s,，。；;、]+", args.query) if term.strip()]

    hits = []
    for index, chunk in enumerate(chunks, start=1):
        score = score_chunk(chunk, terms)
        if score:
            hits.append((score, index, chunk))

    hits.sort(key=lambda item: (-item[0], item[1]))
    for score, index, chunk in hits[: args.limit]:
        print(f"[{index}] score={score}")
        print(chunk)
        print()


def main():
    parser = argparse.ArgumentParser(description="Extract and search OCRed school-rule DOCX files.")
    subparsers = parser.add_subparsers(required=True)

    extract_parser = subparsers.add_parser("extract")
    extract_parser.add_argument("source")
    extract_parser.add_argument("-o", "--output", default="work/docx_text.txt")
    extract_parser.set_defaults(func=extract)

    search_parser = subparsers.add_parser("search")
    search_parser.add_argument("query")
    search_parser.add_argument("-t", "--text", default="work/docx_text.txt")
    search_parser.add_argument("-n", "--limit", type=int, default=8)
    search_parser.set_defaults(func=search)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
